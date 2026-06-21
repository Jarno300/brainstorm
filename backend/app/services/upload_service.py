"""
File upload service — accepts PDF, text, and image files.
Extracts text content and stores files on disk for injection into chat context.
"""

import io
import logging
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from app.config import APP_ENV
from app.models.brainstorm import Brainstorm
from app.services.library_service import create_library_entry

logger = logging.getLogger(__name__)

# ── Configuration ───────────────────────────────────────────
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./uploads"))
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json",
    ".pdf", ".docx",
    ".png", ".jpg", ".jpeg", ".gif", ".webp",
}
ALLOWED_MIMETYPES = {
    "text/plain", "text/markdown", "text/csv", "application/json",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/png", "image/jpeg", "image/gif", "image/webp",
}


def _ensure_upload_dir(brainstorm_id: uuid.UUID) -> Path:
    """Create and return the upload directory for a brainstorm."""
    path = UPLOAD_DIR / str(brainstorm_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        logger.warning("PDF text extraction failed: %s", e)
        return "[PDF text extraction failed — file may be scanned or image-based]"


def _extract_image_description(file_bytes: bytes, filename: str) -> str:
    """Return a placeholder description for images (LLM vision not yet implemented)."""
    return f"[Image: {filename}]"


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX file with heading structure preserved."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs.append("")
                continue
            # Preserve heading levels as markdown
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if style_name.startswith("heading"):
                    level = 1
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 1
                    level = min(level, 4)  # cap at h4
                    paragraphs.append(f"{'#' * level} {text}")
                    continue
            paragraphs.append(text)
        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("")
                paragraphs.extend(rows)
                paragraphs.append("")
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning("DOCX text extraction failed: %s", e)
        return "[DOCX text extraction failed]"


def extract_text(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Extract text content from an uploaded file.

    Returns the extracted text string, or a placeholder if extraction
    is not possible for this file type.
    """
    ext = Path(filename).suffix.lower()

    # Text-based files
    if ext in (".txt", ".md", ".csv", ".json"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="replace")

    # PDF
    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)

    # DOCX
    if ext == ".docx":
        return _extract_docx_text(file_bytes)

    # Images — placeholder for now (vision support requires multimodal models)
    if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
        return _extract_image_description(file_bytes, filename)

    return f"[Unsupported file: {filename}]"


def validate_upload(file: UploadFile) -> None:
    """Validate file size and type before processing."""
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )


def save_upload(
    db: Session,
    brainstorm_id: uuid.UUID,
    file: UploadFile,
    user_id: uuid.UUID,
) -> dict:
    """Save an uploaded file, extract text, and return metadata.

    The extracted text is returned so the caller can inject it into
    the chat context.
    """
    # Verify ownership
    brainstorm = (
        db.query(Brainstorm)
        .filter(Brainstorm.id == brainstorm_id, Brainstorm.user_id == user_id, Brainstorm.deleted_at.is_(None))
        .first()
    )
    if not brainstorm:
        raise HTTPException(status_code=404, detail="Brainstorm not found")

    validate_upload(file)

    # Read file bytes
    file_bytes = file.file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum is {MAX_FILE_SIZE // (1024*1024)} MB")

    # Save to disk
    upload_dir = _ensure_upload_dir(brainstorm_id)
    file_id = uuid.uuid4()
    safe_name = f"{file_id}_{Path(file.filename or 'upload').name}"
    dest = upload_dir / safe_name

    dest.write_bytes(file_bytes)

    # Extract text
    extracted = extract_text(file_bytes, file.filename or "unknown", file.content_type or "")

    logger.info(
        "File uploaded | brainstorm=%s file=%s size=%d chars=%d",
        brainstorm_id, file.filename, len(file_bytes), len(extracted),
    )

    # Persist extracted text as a library entry for searchability and durability
    try:
        folder_name = Path(file.filename or "upload").stem or "Uploads"
        entry = create_library_entry(
            db=db,
            brainstorm_id=brainstorm_id,
            topic_id=None,
            folder_name=folder_name,
            file_name=file.filename or "upload",
            content=extracted,
            commit=True,
            source_type="upload",
            source_id=str(file_id),
        )
    except Exception as lib_err:
        logger.warning("Failed to create library entry for upload: %s", lib_err)

    return {
        "id": str(file_id),
        "filename": file.filename,
        "size": len(file_bytes),
        "text": extracted,
        "path": str(dest),
    }


def get_uploads(brainstorm_id: uuid.UUID) -> List[dict]:
    """List uploaded files for a brainstorm."""
    upload_dir = _ensure_upload_dir(brainstorm_id)
    files = []
    for f in sorted(upload_dir.iterdir(), key=lambda x: x.stat().st_mtime, reverse=True):
        if f.is_file():
            files.append({
                "id": f.stem.split("_")[0],
                "filename": "_".join(f.stem.split("_")[1:]),
                "size": f.stat().st_size,
                "uploaded_at": datetime.fromtimestamp(f.stat().st_mtime, tz=timezone.utc).isoformat(),
            })
    return files
